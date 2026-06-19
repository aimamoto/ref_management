import sys
import re
import argparse
from pathlib import Path
from collections import defaultdict
from typing import Dict, List, Optional, Any, Tuple, Match

# --- MONKEY PATCH FOR PYPARSING/BIBTEXPARSER COMPATIBILITY ---
import pyparsing
if not hasattr(pyparsing, 'DelimitedList'):
    if hasattr(pyparsing, 'delimited_list'): setattr(pyparsing, 'DelimitedList', pyparsing.delimited_list)
    elif hasattr(pyparsing, 'delimitedList'): setattr(pyparsing, 'DelimitedList', pyparsing.delimitedList)

import bibtexparser
from bibtexparser.customization import convert_to_unicode, author, type as type_cust
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from rapidfuzz import fuzz

# --- INTELLIGENCE DICTIONARIES ---
AA_LIST_3 = "Ala|Arg|Asn|Asp|Cys|Gln|Glu|Gly|His|Ile|Leu|Lys|Met|Phe|Pro|Ser|Thr|Trp|Tyr|Val"
AA_PATTERN_3 = re.compile(rf'\b(?:{AA_LIST_3})\s*-?\s*$', re.IGNORECASE)
AA_PATTERN_1 = re.compile(r'\b[ARNDCQEGHILKMFPSTWY]-?$')
MATH_UNIT_LIST = ["CV", "R", "r", "m", "cm", "mm", "µm", "um", "nm", "km", "kg", "x", "y", "z", "p", "n", "k", "v", "V", "D", "Ca", "Mg", "Na", "K", "Cl", "Fe", "Zn", "Cu", "O", "H", "N", "C", "P", "S", "M", "χ", "Χ"]
MATH_UNIT_PATTERN = re.compile(rf'\b(?:{"|".join(MATH_UNIT_LIST)})\s*$') 
POWER_PATTERN = re.compile(r'(?:10|x10|×10|\*10)\s*$')
IGNORE_PREFIXES = re.compile(r'(?i:\b(?:fig(?:ure)?|eq(?:uation)?|tbl|table|section|sec|step)\s*\.?\s*)$')

PROSE_STOP_WORDS = {'the', 'is', 'are', 'was', 'were', 'that', 'this', 'to', 'for', 'with', 
                    'in', 'on', 'by', 'an', 'we', 'our', 'as', 'it', 'can', 'be', 'has', 'have', 'of', 'and', 'from', 'which'}

REF_HEADER_PATTERN = re.compile(r'^\s*(?:[0-9]+\.?\s*)?(?:REFERENCES|BIBLIOGRAPHY|LITERATURE CITED|WORKS CITED)\s*$', re.IGNORECASE)

# --- CITATION MANAGER CLASS ---
class CitationManager:
    def __init__(self, bib_file: Path, style: str):
        self.bib_file = bib_file
        self.style = style
        self.entries: Dict[int, Dict[str, Any]] = {}      
        self.dup_map: Dict[int, int] = {}      
        self.appearance_order: Dict[int, int] = {}
        self.current_seq: int = 1
        self.citations_text: Dict[int, str] = {}    
        self.suffixes: Dict[int, str] = {} 
        self.update_count: int = 0
        self.load_data()

    def clean_text(self, text: Any) -> str:
        if not text: return ""
        text_str = str(text)
        while '{' in text_str or '}' in text_str: text_str = text_str.replace('{', '').replace('}', '')
        return re.sub(r'\s+', ' ', text_str).strip()

    def normalize_page_range(self, page_str: str) -> str:
        return self.clean_text(page_str).replace('--', '–').replace('-', '–').strip()

    def get_ref_number(self, id_str: str) -> Optional[int]:
        match = re.search(r'(\d+)', id_str)
        return int(match.group(1)) if match else None

    def get_citation_base(self, entry: Dict[str, Any]) -> str:
        auths = entry.get('author', [])
        if not auths: return "Unknown"
        clean_auths = [self.clean_text(a.split(',')[0].strip() if ',' in a else a.split()[-1].strip()) for a in auths]
        if len(clean_auths) == 1: return clean_auths[0]
        elif len(clean_auths) == 2: return f"{clean_auths[0]} and {clean_auths[1]}"
        else: return f"{clean_auths[0]} et al."

    def format_authors(self, entry: Dict[str, Any]) -> str:
        auths = entry.get('author', [])
        formatted_auths = []
        for a in auths:
            if ',' in a:
                parts = a.split(',')
                surname = self.clean_text(parts[0])
                initials = "".join([p.strip()[0].upper() for p in parts[1].split() if p.strip()])
            else:
                parts = a.split()
                surname = self.clean_text(parts[-1]) if parts else "Unknown"
                initials = "".join([p.strip()[0].upper() for p in parts[:-1] if p.strip()]) if len(parts) > 1 else ""
            
            if self.style == 'numbered':
                formatted_auths.append(f"{surname} {initials}".strip())
            else:
                formatted_auths.append(f"{surname}, {'.'.join(initials)}." if initials else surname)
        
        if len(formatted_auths) > 1:
            joiner = " and " if self.style == 'numbered' else ", and "
            return ", ".join(formatted_auths[:-1]) + joiner + formatted_auths[-1]
        return formatted_auths[0] if formatted_auths else "Unknown"

    def format_bibliography_entry(self, entry: Dict[str, Any], seq_num: Optional[int] = None) -> str:
        auth_str = self.format_authors(entry)
        ref_num = self.get_ref_number(entry.get('ID', ''))
        year = self.clean_text(entry.get('year', '????'))
        
        if self.style == 'author-year':
            suffix = self.suffixes.get(ref_num, "") if ref_num else ""
            year = f"{year}{suffix}"

        title = self.clean_text(entry.get('title', ''))
        title = re.sub(r'https?://[^\s]+', '', title).strip() 
        if title and title[-1] not in '.?': title += "."
        
        journal_str = self.clean_text(entry.get('journal', ''))
        volume = self.clean_text(entry.get('volume', ''))
        pages = self.normalize_page_range(entry.get('pages', ''))
        
        if volume: journal_str += f" {volume}"
        if pages: journal_str += f", {pages}" if volume else f" {pages}"
        if journal_str and journal_str[-1] not in '.?': journal_str += "."
        
        doi_str = f" doi: {self.clean_text(entry['doi']).replace('https://doi.org/', '')}" if entry.get('doi') else ""
        
        if self.style == 'numbered':
            return f"{seq_num}. {auth_str} ({year}) {title} {journal_str}{doi_str}".strip()
        else:
            return f"{auth_str} ({year}). {title} {journal_str}{doi_str}".strip()

    def load_data(self):
        print(f"Loading BibTeX data from: {self.bib_file}...")
        try:
            with open(self.bib_file, encoding='utf-8') as f:
                parser = bibtexparser.bparser.BibTexParser(common_strings=True)
                parser.customization = lambda r: convert_to_unicode(author(type_cust(r)))
                db = parser.parse_file(f)
        except Exception as e:
            sys.exit(1)

        seen_pmids, seen_dois, seen_titles = {}, {}, {}
        sorted_entries = sorted(db.entries, key=lambda x: self.get_ref_number(x['ID']) or 9999)

        for entry in sorted_entries:
            ref_num = self.get_ref_number(entry['ID'])
            if not ref_num: continue

            canonical_ref = seen_pmids.get(entry.get('pmid')) or seen_dois.get(entry.get('doi'))
            if not canonical_ref:
                title = self.clean_text(entry.get('title', ''))
                norm_title = re.sub(r'[^\w]', '', title.lower())
                if len(norm_title) > 20:
                    for s_title, s_ref in seen_titles.items():
                        if fuzz.ratio(norm_title, s_title) > 95:
                            canonical_ref = s_ref; break

            if canonical_ref: self.dup_map[ref_num] = canonical_ref
            else:
                self.dup_map[ref_num] = ref_num
                if entry.get('pmid'): seen_pmids[entry['pmid']] = ref_num
                if entry.get('doi'): seen_dois[entry['doi']] = ref_num
                norm_title = re.sub(r'[^\w]', '', self.clean_text(entry.get('title', '')).lower())
                if len(norm_title) > 20: seen_titles[norm_title] = ref_num
                self.entries[ref_num] = entry

        if self.style == 'author-year':
            collision_groups = defaultdict(list)
            for ref_num, entry in self.entries.items():
                collision_groups[(self.get_citation_base(entry), self.clean_text(entry.get('year', '????')))].append(ref_num)
            for key, ref_nums in collision_groups.items():
                if len(ref_nums) > 1:
                    ref_nums.sort(key=lambda r: self.entries[r].get('title', '').lower())
                    for i, r_num in enumerate(ref_nums):
                        if i < 26: self.suffixes[r_num] = "abcdefghijklmnopqrstuvwxyz"[i]
            for ref_num, entry in self.entries.items():
                base_auth = self.get_citation_base(entry)
                year = f"{self.clean_text(entry.get('year', '????'))}{self.suffixes.get(ref_num, '')}"
                self.citations_text[ref_num] = f"{base_auth}, {year}"

# --- HELPER & DOCX PROCESSING ---

def format_ranges(nums: List[int]) -> str:
    if not nums: return ""
    nums = sorted(list(set(nums)))
    ranges, start, prev = [], nums[0], nums[0]
    for n in nums[1:]:
        if n == prev + 1: prev = n
        else:
            ranges.append(str(start) if prev == start else (f"{start}, {prev}" if prev == start + 1 else f"{start}–{prev}"))
            start = prev = n
    ranges.append(str(start) if prev == start else (f"{start}, {prev}" if prev == start + 1 else f"{start}–{prev}"))
    return ", ".join(ranges)

def iter_block_items(doc):
    for child in doc.element.body:
        if isinstance(child, CT_P): yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl): yield Table(child, doc)

def replace_text_preserve_formatting(para: Paragraph, pattern: re.Pattern, callback):
    text = para.text
    matches = list(pattern.finditer(text))
    if not matches: return

    run_map = []
    for r_idx, run in enumerate(para.runs):
        for c_idx in range(len(run.text)): run_map.append((r_idx, c_idx))
    
    if len(run_map) != len(text):
        for run in para.runs:
            if pattern.search(run.text): run.text = pattern.sub(callback, run.text)
        return

    for match in reversed(matches):
        start, end = match.span()
        replacement = callback(match)
        if replacement == match.group(0): continue
        
        start_r_idx, start_c_idx = run_map[start]
        end_r_idx, end_c_idx = run_map[end - 1]
        
        if start_r_idx == end_r_idx:
            run = para.runs[start_r_idx]
            run.text = run.text[:start_c_idx] + replacement + run.text[end_c_idx + 1:]
        else:
            run_start = para.runs[start_r_idx]
            run_start.text = run_start.text[:start_c_idx] + replacement
            for r_idx in range(start_r_idx + 1, end_r_idx): para.runs[r_idx].text = ""
            run_end = para.runs[end_r_idx]
            run_end.text = run_end.text[end_c_idx + 1:]

def process_paragraph_content(para: Paragraph, manager: CitationManager, citation_pattern: re.Pattern, in_main_body: bool):
    max_ref = max(manager.entries.keys()) if manager.entries else 0

    preceding_text = ""
    for run in para.runs:
        text = run.text.strip()
        
        # Only process superscripts if we have passed the title page
        if in_main_body and run.font.superscript and re.match(r'^[\d,\s\-–]+$', text):
            
            is_math_power = bool(POWER_PATTERN.search(preceding_text)) and text.isdigit()
            
            if (AA_PATTERN_3.search(preceding_text) or 
                AA_PATTERN_1.search(preceding_text) or 
                MATH_UNIT_PATTERN.search(preceding_text) or 
                IGNORE_PREFIXES.search(preceding_text) or 
                is_math_power):
                
                preceding_text += run.text
                continue
            
            clean_text = text.replace('–', '-')
            is_valid = True
            for part in clean_text.split(','):
                if '-' in part:
                    b = part.split('-')
                    if not (len(b) == 2 and b[0].strip().isdigit() and b[1].strip().isdigit()): is_valid = False
                elif not part.strip().isdigit(): is_valid = False
            
            if is_valid:
                run.font.superscript = False
                run.text = f"[{text}]"
                
        preceding_text += run.text
    
    artifact_pattern = re.compile(r'(?:geometry|ref|source)\.(\d+)', re.IGNORECASE)
    replace_text_preserve_formatting(para, artifact_pattern, lambda m: f"[{m.group(1)}]")

    def replace_callback(match: Match) -> str:
        preceding = para.text[:match.start()]
        
        if (AA_PATTERN_3.search(preceding) or 
            AA_PATTERN_1.search(preceding) or 
            MATH_UNIT_PATTERN.search(preceding) or 
            IGNORE_PREFIXES.search(preceding)):
            return match.group(0)

        raw_inner = match.group(1).replace('–', '-')
        if match.group(0).startswith('(') and raw_inner.isdigit() and 1900 <= int(raw_inner) <= 2100:
            return match.group(0)

        oids = []
        for part in raw_inner.split(','):
            part = part.strip()
            if '-' in part:
                bounds = part.split('-')
                if len(bounds) == 2 and bounds[0].strip().isdigit() and bounds[1].strip().isdigit():
                    start, end = int(bounds[0].strip()), int(bounds[1].strip())
                    if start <= end and (end - start) < 50:
                        oids.extend(range(start, end + 1))
                    else: return match.group(0)
                else: return match.group(0)
            else:
                if not part.isdigit(): return match.group(0)
                oids.append(int(part))
        
        if not any(manager.dup_map.get(o, o) in manager.entries for o in oids):
            if all(o > max_ref + 20 for o in oids): return match.group(0)

        new_citations = []
        manager.update_count += 1 # Track successful updates
        
        for oid in oids:
            cid = manager.dup_map.get(oid, oid)
            if cid in manager.entries:
                if manager.style == 'numbered':
                    if cid not in manager.appearance_order:
                        manager.appearance_order[cid] = manager.current_seq
                        manager.current_seq += 1
                    new_citations.append(manager.appearance_order[cid])
                else:
                    new_citations.append(manager.citations_text[cid])
            else:
                new_citations.append(f"!!MISSING[{oid}]!!" if manager.style == 'numbered' else f"(!!MISSING[{oid}]!!)")
        
        if manager.style == 'numbered':
            ints = [c for c in new_citations if isinstance(c, int)]
            strs = [c for c in new_citations if isinstance(c, str)]
            range_str = format_ranges(ints)
            return f"[{', '.join(([range_str] if range_str else []) + strs)}]"
        else:
            clean = [c.strip('()') for c in new_citations]
            return f"({'; '.join(clean)})"

    replace_text_preserve_formatting(para, citation_pattern, replace_callback)

    if manager.style == 'numbered':
        punctuation_pattern = re.compile(r'([^\.\s])\.\s*(\[[\d,\s–\-]+\])')
    else:
        punctuation_pattern = re.compile(r'([^\.\s])\.\s*(\([A-Z][^)]+\d{4}[a-z]?\))')
    replace_text_preserve_formatting(para, punctuation_pattern, lambda m: f"{m.group(1)} {m.group(2)}.")

def process_document(docx_path: Path, output_path: Path, manager: CitationManager):
    print(f"\nProcessing document: {docx_path} (Style: {manager.style.upper()})...")
    doc = Document(str(docx_path))
    citation_pattern = re.compile(r'[\[\(]([\d\s,\-–]+)[\]\)]')

    main_font = None
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name:
                main_font = r.font.name
                break
        if main_font: break

    ref_header_element = None
    for p in doc.paragraphs:
        if REF_HEADER_PATTERN.match(p.text):
            ref_header_element = p._element
            break

    in_main_body = False 
    block_counter = 0

    for block in iter_block_items(doc):
        block_counter += 1
        
        if isinstance(block, Paragraph):
            if ref_header_element is not None and block._element == ref_header_element: break 

            if not in_main_body:
                text_clean = block.text.strip().lower()
                if text_clean in ['abstract', 'introduction', 'background', 'summary', 'methods', 'results']:
                    in_main_body = True
                    print("   -> [Info] Main text body detected. Superscript scanning activated.")
                else:
                    words = set(re.findall(r'\b[a-z]+\b', text_clean))
                    if len(words) >= 15 and len(words.intersection(PROSE_STOP_WORDS)) >= 4:
                        in_main_body = True
                        print("   -> [Info] Main text body detected. Superscript scanning activated.")
                
                # Failsafe: Turn off shield automatically after 25 blocks
                if block_counter > 25 and not in_main_body:
                    in_main_body = True
                    print("   -> [Info] Title Page Shield auto-disabled (reached block 25).")

            # Always process, but pass the `in_main_body` state to protect superscripts
            process_paragraph_content(block, manager, citation_pattern, in_main_body)
            
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph_content(para, manager, citation_pattern, in_main_body)

    ref_header_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p._element == ref_header_element:
            ref_header_index = i
            break
    
    def insert_reference_heading():
        p = doc.add_paragraph()
        run = p.add_run('References')
        run.bold = True
        if main_font: run.font.name = main_font

    if ref_header_index != -1:
        paragraphs_to_remove = doc.paragraphs[ref_header_index:]
        for p in paragraphs_to_remove:
            p_elm = p._element
            parent = p_elm.getparent()
            if parent is not None:
                parent.remove(p_elm)
        insert_reference_heading()
    else:
        doc.add_page_break()
        insert_reference_heading()

    if manager.style == 'numbered':
        valid_entries = [num for num in manager.entries if num in manager.appearance_order]
        valid_entries.sort(key=lambda num: manager.appearance_order[num])
        for num in valid_entries:
            text = manager.format_bibliography_entry(manager.entries[num], manager.appearance_order[num])
            p = doc.add_paragraph()
            run = p.add_run(text)
            if main_font: run.font.name = main_font 
    else:
        valid_entries = [manager.entries[num] for num in manager.entries if num in manager.citations_text]
        valid_entries.sort(key=lambda e: (e.get('author', ["z"])[0] if isinstance(e.get('author'), list) else str(e.get('author', "z"))).lower())
        for entry in valid_entries:
            text = manager.format_bibliography_entry(entry)
            p = doc.add_paragraph()
            run = p.add_run(text)
            if main_font: run.font.name = main_font 

    doc.save(str(output_path))
    print(f"Success! Saved to {output_path}")
    print(f" -> Tracked and updated {manager.update_count} in-text citations.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("bib", type=Path, help="Verified .bib file")
    parser.add_argument("doc", type=Path, help="Input .docx file")
    parser.add_argument("--style", choices=['1', '2', 'numbered', 'author-year'], help="1: Numbered, 2: Author-Year")
    args = parser.parse_args()
    
    style_choice = args.style
    if not style_choice:
        print("\n=== Select Output Citation Style ===")
        print("  1. Sequential Numbered (e.g., [1-3] -> 1. Lopez-Otin C...)")
        print("  2. Author-Year (e.g., (Smith, 2023) -> Smith, I. (2023)...)")
        while True:
            choice = input("Enter 1 or 2: ").strip()
            if choice in ['1', '2']:
                style_choice = choice
                break
    
    active_style = 'author-year' if style_choice in ['2', 'author-year'] else 'numbered'
    output = args.doc.with_name(f"{args.doc.stem}_final_{active_style}.docx")
    
    mgr = CitationManager(args.bib, style=active_style)
    process_document(args.doc, output, mgr)
