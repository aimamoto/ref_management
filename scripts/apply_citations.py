import bibtexparser
from bibtexparser.customization import convert_to_unicode, author, type as type_cust
import re
import argparse
import os
import sys
from docx import Document
from rapidfuzz import fuzz
from collections import defaultdict

class CitationManager:
    def __init__(self, bib_file):
        self.bib_file = bib_file
        self.entries = {}      
        self.dup_map = {}      
        self.citations = {}    
        self.suffixes = {} 
        self.load_data()

    def clean_text(self, text):
        if not text: return ""
        text = str(text)
        while '{' in text or '}' in text:
            text = text.replace('{', '').replace('}', '')
        return re.sub(r'\s+', ' ', text).strip()

    def normalize_page_range(self, page_str):
        """
        Converts '137-42' -> '137-142' and '137--142' -> '137-142'.
        """
        # 1. Standardize dash
        page_str = self.clean_text(page_str).replace('--', '-').strip()
        if '-' not in page_str: return page_str
        
        # 2. Split
        parts = page_str.split('-')
        if len(parts) != 2: return page_str
        
        start, end = parts[0].strip(), parts[1].strip()
        
        # 3. Numeric Expansion Logic
        # Only applies if both sides are purely digits
        if start.isdigit() and end.isdigit():
            # If end is shorter than start (e.g., 42 is shorter than 137)
            if len(end) < len(start):
                # Calculate prefix from start
                # e.g. Start=137, End=42. Missing 1 char. Prefix = '1'
                prefix = start[:len(start)-len(end)]
                expanded_end = prefix + end
                
                # Logical check: Expanded end must be greater than start
                if int(expanded_end) > int(start):
                    return f"{start}-{expanded_end}"
        
        return f"{start}-{end}"

    def get_ref_number(self, id_str):
        match = re.search(r'(\d+)', id_str)
        return int(match.group(1)) if match else None

    def get_citation_base(self, entry):
        auths = entry.get('author', [])
        if not auths: return "Unknown"
        clean_auths = []
        for a in auths:
            if ',' in a: surname = a.split(',')[0].strip()
            else: surname = a.split()[-1].strip()
            clean_auths.append(self.clean_text(surname))

        if len(clean_auths) == 1: return clean_auths[0]
        elif len(clean_auths) == 2: return f"{clean_auths[0]} and {clean_auths[1]}"
        else: return f"{clean_auths[0]} et al."

    def format_citation_text(self, entry, ref_num):
        title = self.clean_text(entry.get('title', '')).upper()
        if title in ['REVIEWS', 'UNKNOWN', 'REFERENCES'] and 'author' not in entry:
            return None 

        year = self.clean_text(entry.get('year', '????'))
        suffix = self.suffixes.get(ref_num, "")
        full_year = f"{year}{suffix}"
        base_auth = self.get_citation_base(entry)
        
        return f"({base_auth}, {full_year})"

    def load_data(self):
        print(f"Loading BibTeX data from: {self.bib_file}...")
        if self.bib_file.endswith('.txt'):
            print("\n!!! ERROR: Use the .bib file, not .txt !!!")
            sys.exit(1)

        with open(self.bib_file, encoding='utf-8') as f:
            parser = bibtexparser.bparser.BibTexParser(common_strings=True)
            parser.customization = lambda r: convert_to_unicode(author(type_cust(r)))
            db = parser.parse_file(f)

        if not db.entries:
            print("\n!!! ERROR: No references found !!!")
            sys.exit(1)

        print(f"   > Parsed {len(db.entries)} references.")

        seen_pmids = {}
        seen_dois = {}
        seen_titles = {} 
        sorted_entries = sorted(db.entries, key=lambda x: self.get_ref_number(x['ID']) or 9999)

        for entry in sorted_entries:
            ref_num = self.get_ref_number(entry['ID'])
            if not ref_num: continue

            canonical_ref = None
            pmid = entry.get('pmid')
            doi = entry.get('doi')
            
            if pmid and pmid in seen_pmids: canonical_ref = seen_pmids[pmid]
            elif doi and doi in seen_dois: canonical_ref = seen_dois[doi]

            if not canonical_ref:
                title = self.clean_text(entry.get('title', ''))
                norm_title = re.sub(r'[^\w]', '', title.lower())
                if len(norm_title) > 20:
                    for s_title, s_ref in seen_titles.items():
                        if fuzz.ratio(norm_title, s_title) > 95:
                            canonical_ref = s_ref
                            print(f"   > Fuzzy Duplicate: Ref{ref_num} maps to Ref{canonical_ref}")
                            break

            if canonical_ref:
                self.dup_map[ref_num] = canonical_ref
            else:
                self.dup_map[ref_num] = ref_num
                if pmid: seen_pmids[pmid] = ref_num
                if doi: seen_dois[doi] = ref_num
                title = self.clean_text(entry.get('title', ''))
                norm_title = re.sub(r'[^\w]', '', title.lower())
                if len(norm_title) > 20: seen_titles[norm_title] = ref_num
                self.entries[ref_num] = entry

        # Disambiguation (2025a vs 2025b)
        collision_groups = defaultdict(list)
        for ref_num, entry in self.entries.items():
            base_auth = self.get_citation_base(entry)
            year = self.clean_text(entry.get('year', '????'))
            key = (base_auth, year)
            collision_groups[key].append(ref_num)
        
        for key, ref_nums in collision_groups.items():
            if len(ref_nums) > 1:
                ref_nums.sort(key=lambda r: self.entries[r].get('title', '').lower())
                chars = "abcdefghijklmnopqrstuvwxyz"
                for i, r_num in enumerate(ref_nums):
                    if i < len(chars):
                        self.suffixes[r_num] = chars[i]

        for ref_num, entry in self.entries.items():
            self.citations[ref_num] = self.format_citation_text(entry, ref_num)

# --- DOCX PROCESSING ---

def process_paragraph_content(para, manager, citation_pattern):
    # 1. Superscripts
    for run in para.runs:
        text = run.text.strip()
        if run.font.superscript and re.match(r'^[\d,]+$', text):
            run.font.superscript = False
            run.text = f"[{text}]"
    
    # 2. Artifacts
    artifact_pattern = re.compile(r'(?:geometry|ref|source)\.(\d+)', re.IGNORECASE)
    if artifact_pattern.search(para.text):
        para.text = artifact_pattern.sub(r'[\1]', para.text)

    # 3. Replace Citations
    if citation_pattern.search(para.text):
        def replace_callback(match):
            content = match.group(1)
            raw_nums = [n.strip() for n in content.split(',')]
            new_citations = []
            for n in raw_nums:
                if not n.isdigit(): return match.group(0)
                oid = int(n)
                cid = manager.dup_map.get(oid, oid)
                if cid in manager.citations:
                    txt = manager.citations[cid]
                    new_citations.append(txt if txt else f"(!!UNKNOWN Ref[{oid}]!!)")
                else:
                    new_citations.append(f"(!!MISSING_REF[{oid}]!!)")
            clean = [c.strip('()') for c in new_citations]
            return f"({'; '.join(clean)})"
        para.text = citation_pattern.sub(replace_callback, para.text)

    # 4. Fix Punctuation
    punctuation_pattern = re.compile(r'([^\.\s])\.\s*(\([A-Z][^)]+\d{4}[a-z]?\))')
    if punctuation_pattern.search(para.text):
        para.text = punctuation_pattern.sub(r'\1 \2.', para.text)

def process_document(docx_path, output_path, manager):
    print(f"Processing document: {docx_path}...")
    doc = Document(docx_path)
    citation_pattern = re.compile(r'[\[\(]([\d\s,]+)[\]\)]')

    print("   > Scanning Body Text...")
    for para in doc.paragraphs:
        process_paragraph_content(para, manager, citation_pattern)

    print("   > Scanning Tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph_content(para, manager, citation_pattern)

    # Rebuild Bibliography
    ref_header_index = -1
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() in ['references', 'works cited', 'bibliography']:
            ref_header_index = i
            break
    
    if ref_header_index != -1:
        print("Removing old References section...")
        for i in range(ref_header_index + 1, len(doc.paragraphs)):
            doc.paragraphs[i].clear()
    else:
        print("Appending new References section...")
        doc.add_page_break()
        doc.add_heading('References', level=1)

    valid_entries = [e for num, e in manager.entries.items() if manager.citations.get(num)]
    valid_entries.sort(key=lambda e: (e.get('author', [])[0] if e.get('author') else "z").lower())

    for entry in valid_entries:
        doc.add_paragraph(format_bib_entry(entry, manager))

    doc.save(output_path)
    print(f"Success! Saved to {output_path}")

def format_bib_entry(entry, manager):
    ref_num = manager.get_ref_number(entry.get('ID', ''))
    def clean(t): return manager.clean_text(t)
    
    auths = entry.get('author', [])
    formatted_auths = []
    for a in auths:
        if ',' in a: parts = a.split(',')
        else: parts = a.split() 
        surname = clean(parts[0])
        initial = clean(parts[1])[0] if len(parts) > 1 else ""
        formatted_auths.append(f"{surname}, {initial}.")
    
    if len(formatted_auths) > 1: auth_str = ", ".join(formatted_auths[:-1]) + ", and " + formatted_auths[-1]
    else: auth_str = formatted_auths[0] if formatted_auths else "Unknown"

    year = clean(entry.get('year', '????'))
    suffix = manager.suffixes.get(ref_num, "")
    full_year = f"{year}{suffix}"

    title = clean(entry.get('title', ''))
    if title and title[-1] not in '.?': title += "."
    
    journal = clean(entry.get('journal', ''))
    volume = clean(entry.get('volume', ''))
    
    # --- Page Normalization ---
    pages_raw = clean(entry.get('pages', ''))
    pages = manager.normalize_page_range(pages_raw)
    
    journal_str = f"{journal}"
    if volume: journal_str += f" {volume}"
    if pages: journal_str += f", {pages}"
    
    ids = []
    if entry.get('pmid'): ids.append(f"PMID: {clean(entry['pmid'])}")
    if entry.get('doi'): ids.append(f"doi: {clean(entry['doi']).replace('https://doi.org/','')}")
    
    return f"{auth_str} ({full_year}). {title} {journal_str} {' '.join(ids)}".strip()

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("bib", help="Verified .bib file")
    parser.add_argument("doc", help="Input .docx file")
    args = parser.parse_args()
    
    base, ext = os.path.splitext(args.doc)
    output = f"{base}_final.docx"
    
    mgr = CitationManager(args.bib)
    process_document(args.doc, output, mgr)
