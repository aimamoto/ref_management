import sys
import re
import argparse
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Match
import html

# --- MONKEY PATCH FOR PYPARSING/BIBTEXPARSER COMPATIBILITY ---
import pyparsing
if not hasattr(pyparsing, 'DelimitedList'):
    if hasattr(pyparsing, 'delimited_list'): setattr(pyparsing, 'DelimitedList', pyparsing.delimited_list)
    elif hasattr(pyparsing, 'delimitedList'): setattr(pyparsing, 'DelimitedList', pyparsing.delimitedList)

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

# --- CITEPROC IMPORTS ---
from citeproc import CitationStylesStyle, CitationStylesBibliography
from citeproc import Citation, CitationItem
from citeproc import formatter
from citeproc.source.bibtex import BibTeX

# --- INTELLIGENCE DICTIONARIES (SMART SHIELDS PRESERVED) ---
AA_LIST_3 = "Ala|Arg|Asn|Asp|Cys|Gln|Glu|Gly|His|Ile|Leu|Lys|Met|Phe|Pro|Ser|Thr|Trp|Tyr|Val"
AA_PATTERN_3 = re.compile(rf'\b(?:{AA_LIST_3})\s*-?\s*$', re.IGNORECASE)
AA_PATTERN_1 = re.compile(r'\b[ARNDCQEGHILKMFPSTWY]-?$')
MATH_UNIT_LIST = ["CV", "R", "r", "m", "cm", "mm", "µm", "um", "nm", "km", "kg", "x", "y", "z", "p", "n", "k", "v", "V", "D", "Ca", "Mg", "Na", "K", "Cl", "Fe", "Zn", "Cu", "O", "H", "N", "C", "P", "S", "M", "χ", "Χ"]
MATH_UNIT_PATTERN = re.compile(rf'\b(?:{"|".join(MATH_UNIT_LIST)})\s*$') 
POWER_PATTERN = re.compile(r'(?:10|x10|×10|\*10)\s*$')
IGNORE_PREFIXES = re.compile(r'(?i:\b(?:fig(?:ure)?|eq(?:uation)?|tbl|table|section|sec|step)\s*\.?\s*)$')

PROSE_STOP_WORDS = {'the', 'is', 'are', 'was', 'were', 'that', 'this', 'to', 'for', 'with', 
                    'in', 'on', 'by', 'an', 'we', 'our', 'as', 'it', 'can', 'be', 'has', 'have', 'of', 'and', 'from', 'which'}

REF_HEADER_PATTERN = re.compile(r'^\s*(?:[0-9]+\.?\s*)?(?:REFERENCES|BIBLIOGRAPHY|LITERATURE CITED|WORKS CITED)\s*$', re.IGNORECASE)

# --- CITATION MANAGER CLASS (NOW POWERED BY CSL) ---
class CSLCitationManager:
    def __init__(self, bib_file: Path, csl_file: Path):
        self.bib_file = bib_file
        self.csl_file = csl_file
        self.update_count: int = 0
        
        # Load CSL Engine
        print(f"Loading Bibliography Data and CSL style ({csl_file.name})...")
        # Ensure utf-8 encoding is explicitly passed to prevent ascii decode errors
        self.bib_source = BibTeX(str(self.bib_file), encoding='utf-8')
        self.bib_style = CitationStylesStyle(str(self.csl_file))
        self.bibliography = CitationStylesBibliography(self.bib_style, self.bib_source, formatter.html)
        
        # Map integer reference numbers to Citeproc IDs (usually the bibtex key)
        self.id_map: Dict[int, str] = {}
        for key, entry in self.bib_source.items():
            match = re.search(r'(\d+)', key)
            if match:
                ref_num = int(match.group(1))
                self.id_map[ref_num] = key

    def get_in_text_citation(self, oids: List[int]) -> str:
        """Registers and formats the in-text citation using Citeproc."""
        valid_keys = [self.id_map[o] for o in oids if o in self.id_map]
        
        if not valid_keys:
            return f"[!!MISSING{oids}!!]"

        # Create Citation Items and register them with Citeproc
        citation_items = [CitationItem(key) for key in valid_keys]
        citation = Citation(citation_items)
        self.bibliography.register(citation)
        
        # We output plain text for the in-text citations 
        formatted_cite = self.bibliography.cite(citation, formatter.plain)
        formatted_str = str(formatted_cite)
        
        # --- Post-process to fix citeproc-py collapsing limitations ---
        # If the output is a purely numeric comma-separated list, collapse sequentially (e.g., [1, 2, 3] -> [1–3])
        match = re.match(r'^([\[\(]?)([\d\s,]+)([\]\)]?)$', formatted_str.strip())
        if match:
            prefix, numbers_str, suffix = match.groups()
            nums_raw = [n.strip() for n in numbers_str.split(',')]
            
            # Ensure all split elements are purely digits
            if all(n.isdigit() for n in nums_raw if n):
                nums = sorted(list(set(int(n) for n in nums_raw if n)))
                if nums:
                    ranges, start, prev = [], nums[0], nums[0]
                    for n in nums[1:]:
                        if n == prev + 1:
                            prev = n
                        else:
                            # Collapse logic: 1,2 remains 1,2; 1,2,3 becomes 1–3
                            ranges.append(str(start) if prev == start else (f"{start}, {prev}" if prev == start + 1 else f"{start}–{prev}"))
                            start = prev = n
                    ranges.append(str(start) if prev == start else (f"{start}, {prev}" if prev == start + 1 else f"{start}–{prev}"))
                    
                    # Preserve the original spacing convention (e.g. "[1,2]" vs "[1, 2]")
                    delimiter = ", " if ", " in numbers_str else ","
                    formatted_str = f"{prefix}{delimiter.join(ranges)}{suffix}"

        return formatted_str

# --- HELPER & DOCX PROCESSING ---

def iter_block_items(doc):
    for child in doc.element.body:
        if isinstance(child, CT_P): yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl): yield Table(child, doc)

def count_bracketed_citations(doc) -> int:
    count = 0
    pattern = re.compile(r'[\[\(]([\d\s,\-–]+)[\]\)]')
    
    def check_text(text):
        nonlocal count
        for match in pattern.finditer(text):
            raw_inner = match.group(1).replace('–', '-')
            if match.group(0).startswith('(') and raw_inner.isdigit() and 1900 <= int(raw_inner) <= 2100:
                continue
            parts = raw_inner.split(',')
            valid = True
            for p in parts:
                p = p.strip()
                if '-' in p:
                    b = p.split('-')
                    if not (len(b) == 2 and b[0].strip().isdigit() and b[1].strip().isdigit()): valid = False
                elif not p.isdigit(): valid = False
            if valid: count += 1

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph): check_text(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: check_text(p.text)
    return count

def replace_text_preserve_formatting(para: Paragraph, pattern: re.Pattern, callback):
    text = para.text
    matches = list(pattern.finditer(text))
    if not matches: return

    replacements = [callback(m) for m in matches]

    run_map = []
    for r_idx, run in enumerate(para.runs):
        for c_idx in range(len(run.text)): run_map.append((r_idx, c_idx))
    
    if len(run_map) != len(text):
        new_text = text
        for match, rep in zip(reversed(matches), reversed(replacements)):
            start, end = match.span()
            new_text = new_text[:start] + rep + new_text[end:]
        para.text = new_text
        return

    for match, rep in zip(reversed(matches), reversed(replacements)):
        start, end = match.span()
        if rep == match.group(0): continue
        
        start_r_idx, start_c_idx = run_map[start]
        end_r_idx, end_c_idx = run_map[end - 1]
        
        if start_r_idx == end_r_idx:
            run = para.runs[start_r_idx]
            run.text = run.text[:start_c_idx] + rep + run.text[end_c_idx + 1:]
        else:
            run_start = para.runs[start_r_idx]
            run_start.text = run_start.text[:start_c_idx] + rep
            for r_idx in range(start_r_idx + 1, end_r_idx): para.runs[r_idx].text = ""
            run_end = para.runs[end_r_idx]
            run_end.text = run_end.text[end_c_idx + 1:]

def process_paragraph_content(para: Paragraph, manager: CSLCitationManager, citation_pattern: re.Pattern, in_main_body: bool, convert_superscripts: bool):
    max_ref = max(manager.id_map.keys()) if manager.id_map else 0

    preceding_text = ""
    for run in para.runs:
        text = run.text.strip()
        
        if convert_superscripts and in_main_body and run.font.superscript and re.match(r'^[\d,\s\-–]+$', text):
            
            is_math_power = bool(POWER_PATTERN.search(preceding_text)) and text.isdigit()
            
            if (AA_PATTERN_3.search(preceding_text) or 
                AA_PATTERN_1.search(preceding_text) or 
                MATH_UNIT_PATTERN.search(preceding_text) or 
                IGNORE_PREFIXES.search(preceding_text) or 
                is_math_power):
                
                preceding_text += run.text
                continue
            
            clean_text = text.replace('–', '-')
            is_valid = True
            for part in clean_text.split(','):
                if '-' in part:
                    b = part.split('-')
                    if not (len(b) == 2 and b[0].strip().isdigit() and b[1].strip().isdigit()): is_valid = False
                elif not part.strip().isdigit(): is_valid = False
            
            if is_valid:
                run.font.superscript = False
                run.text = f"[{text}]"
                
        preceding_text += run.text
    
    artifact_pattern = re.compile(r'(?:geometry|ref|source)\.(\d+)', re.IGNORECASE)
    replace_text_preserve_formatting(para, artifact_pattern, lambda m: f"[{m.group(1)}]")

    def replace_callback(match: Match) -> str:
        preceding = para.text[:match.start()]
        
        if (AA_PATTERN_3.search(preceding) or 
            AA_PATTERN_1.search(preceding) or 
            MATH_UNIT_PATTERN.search(preceding) or 
            IGNORE_PREFIXES.search(preceding)):
            return match.group(0)

        raw_inner = match.group(1).replace('–', '-')
        if match.group(0).startswith('(') and raw_inner.isdigit() and 1900 <= int(raw_inner) <= 2100:
            return match.group(0)

        oids = []
        for part in raw_inner.split(','):
            part = part.strip()
            if '-' in part:
                bounds = part.split('-')
                if len(bounds) == 2 and bounds[0].strip().isdigit() and bounds[1].strip().isdigit():
                    start, end = int(bounds[0].strip()), int(bounds[1].strip())
                    if start <= end and (end - start) < 50:
                        oids.extend(range(start, end + 1))
                    else: return match.group(0)
                else: return match.group(0)
            else:
                if not part.isdigit(): return match.group(0)
                oids.append(int(part))
        
        # Execute citation generation
        manager.update_count += 1 
        return manager.get_in_text_citation(oids)

    replace_text_preserve_formatting(para, citation_pattern, replace_callback)

def write_rich_bibliography_entry(doc: Document, html_text: str, main_font: Optional[str]):
    """Parses HTML output from citeproc-py and maps directly to python-docx Rich Text Runs."""
    p = doc.add_paragraph()
    
    # State flags
    is_bold = False
    is_italic = False
    is_smallcaps = False
    is_sup = False
    is_sub = False

    # Regex splits by XML/HTML tags
    tokens = re.split(r'(<[^>]+>)', html_text)
    
    for token in tokens:
        if not token: continue
        
        token_lower = token.lower()
        if token_lower.startswith('<'):
            # Handle styles
            if token_lower in ['<b>', '<strong>']: is_bold = True
            elif token_lower in ['</b>', '</strong>']: is_bold = False
            elif token_lower in ['<i>', '<em>']: is_italic = True
            elif token_lower in ['</i>', '</em>']: is_italic = False
            elif token_lower == '<sup>': is_sup = True
            elif token_lower == '</sup>': is_sup = False
            elif token_lower == '<sub>': is_sub = True
            elif token_lower == '</sub>': is_sub = False
            elif 'small-caps' in token_lower and not token_lower.startswith('</'): is_smallcaps = True
            elif token_lower.startswith('</span'): is_smallcaps = False
            # Ignore structural divs like <div class="csl-entry">
            continue
        
        # It's a text node. Decode HTML entities (e.g., &amp; -> &)
        text_content = html.unescape(token)
        if text_content:
            run = p.add_run(text_content)
            run.bold = is_bold
            run.italic = is_italic
            if is_smallcaps: run.font.small_caps = True
            if is_sup: run.font.superscript = True
            if is_sub: run.font.subscript = True
            if main_font: run.font.name = main_font


def process_document(docx_path: Path, output_path: Path, manager: CSLCitationManager):
    print(f"\nProcessing document: {docx_path} (Using Style: {manager.csl_file.name})...")
    doc = Document(str(docx_path))
    citation_pattern = re.compile(r'[\[\(]([\d\s,\-–]+)[\]\)]')

    main_font = None
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name:
                main_font = r.font.name
                break
        if main_font: break

    bracket_count = count_bracketed_citations(doc)
    # Autodetect if document is using bracketed or superscript styling
    convert_superscripts = (bracket_count < 3)
    
    if convert_superscripts:
        print("   -> [Info] Auto-detected Superscript format. Superscript Converter ENABLED.")
    else:
        print(f"   -> [Info] Auto-detected Bracketed format ({bracket_count} brackets). Superscript Converter DISABLED.")

    ref_header_element = None
    for p in doc.paragraphs:
        if REF_HEADER_PATTERN.match(p.text):
            ref_header_element = p._element
            break

    in_main_body = False 
    block_counter = 0

    for block in iter_block_items(doc):
        block_counter += 1
        if isinstance(block, Paragraph):
            if ref_header_element is not None and block._element == ref_header_element: break 

            if not in_main_body:
                text_clean = block.text.strip().lower()
                if text_clean in ['abstract', 'introduction', 'background', 'summary', 'methods', 'results']:
                    in_main_body = True
                else:
                    words = re.findall(r'\b[a-z]+\b', text_clean)
                    stop_matches = [w for w in words if w in PROSE_STOP_WORDS]
                    if len(words) >= 25 and len(stop_matches) >= 5:
                        in_main_body = True
                
                if block_counter > 25 and not in_main_body:
                    in_main_body = True

            process_paragraph_content(block, manager, citation_pattern, in_main_body, convert_superscripts)
            
        elif isinstance(block, Table):
            if not in_main_body: continue 
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph_content(para, manager, citation_pattern, in_main_body, convert_superscripts)

    # 1. Clear Old Bibliography Header and Content
    ref_header_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p._element == ref_header_element:
            ref_header_index = i
            break
    
    def insert_reference_heading():
        p = doc.add_paragraph()
        run = p.add_run('References')
        run.bold = True
        if main_font: run.font.name = main_font

    if ref_header_index != -1:
        paragraphs_to_remove = doc.paragraphs[ref_header_index:]
        for p in paragraphs_to_remove:
            p_elm = p._element
            parent = p_elm.getparent()
            if parent is not None:
                parent.remove(p_elm)
        insert_reference_heading()
    else:
        doc.add_page_break()
        insert_reference_heading()

    # 2. Rebuild the Unified Bibliography via CSL Engine
    print(f"   -> Rebuilding Rich-Text Bibliography via CSL...")
    manager.bibliography.sort()
    for entry in manager.bibliography.bibliography():
        html_entry = str(entry)
        write_rich_bibliography_entry(doc, html_entry, main_font)

    doc.save(str(output_path))
    print(f"Success! Saved to {output_path}")
    print(f" -> Tracked and dynamically updated {manager.update_count} in-text citation brackets.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("bib", type=Path, help="Verified .bib file")
    parser.add_argument("doc", type=Path, help="Input .docx file")
    parser.add_argument("--csl", type=Path, required=True, help="Path to the CSL style file")
    args = parser.parse_args()
    
    # --- NEW: CSL Path Resolution Logic ---
    csl_path = args.csl
    default_csl_dir = Path("~/citation_styles").expanduser()
    
    if not csl_path.exists():
        alt_path = default_csl_dir / csl_path.name
        if alt_path.exists():
            csl_path = alt_path
        elif not csl_path.suffix == '.csl':
            alt_path_ext = default_csl_dir / f"{csl_path.name}.csl"
            if alt_path_ext.exists():
                csl_path = alt_path_ext
                
    if not csl_path.exists():
        print(f"❌ ERROR: CSL file '{args.csl}' not found locally or in {default_csl_dir}.")
        sys.exit(1)
    # ---------------------------------------
        
    output = args.doc.with_name(f"{args.doc.stem}_final_{csl_path.stem}.docx")
    
    mgr = CSLCitationManager(args.bib, csl_path)
    process_document(args.doc, output, mgr)
