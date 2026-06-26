import sys
import re
import argparse
from pathlib import Path

# --- MONKEY PATCH FOR PYPARSING/BIBTEXPARSER COMPATIBILITY ---
import pyparsing
if not hasattr(pyparsing, 'DelimitedList'):
    if hasattr(pyparsing, 'delimited_list'): setattr(pyparsing, 'DelimitedList', pyparsing.delimited_list)
    elif hasattr(pyparsing, 'delimitedList'): setattr(pyparsing, 'DelimitedList', pyparsing.delimitedList)

import bibtexparser
from docx import Document
from rapidfuzz import fuzz

REF_HEADER_PATTERN = re.compile(r'^\s*(?:[0-9]+\.?\s*)?(?:REFERENCES|BIBLIOGRAPHY|LITERATURE CITED|WORKS CITED)\s*$', re.IGNORECASE)
POST_REF_PATTERN = re.compile(r'^\s*(?:Tables?|Figures?|Figure Legends?|Supplementary.*?|Appendices|Data Availability|Acknowledgements?|Author Contributions?|Funding|Conflict(?:s)? of Interest|Competing Interests?|(?:Table|Figure|Fig\.?)\s*\d+.*)$', re.IGNORECASE)

def clean_for_match(text: str) -> str:
    """Removes punctuation and normalizes spacing for accurate fuzzy matching."""
    if not text: return ""
    text = text.replace('{', '').replace('}', '')
    return re.sub(r'[^\w\s]', '', text.lower()).strip()

def process_document(bib_path: Path, docx_path: Path, output_path: Path):
    print(f"\nReading verified BibTeX: {bib_path.name}...")
    try:
        with open(bib_path, 'r', encoding='utf-8') as f:
            bib_db = bibtexparser.load(f)
    except Exception as e:
        print(f"❌ ERROR reading BibTeX: {e}")
        sys.exit(1)

    # Build an index of cleaned titles to DOIs
    doi_map = {}
    for entry in bib_db.entries:
        doi = entry.get('doi', '').strip()
        title = entry.get('title', '').strip()
        if doi and title:
            # Clean DOI prefix if present
            clean_doi = doi.replace('https://doi.org/', '').replace('doi:', '').strip()
            doi_map[clean_for_match(title)] = clean_doi

    print(f"Loaded {len(doi_map)} DOIs from BibTeX.")
    print(f"Scanning document: {docx_path.name}...")
    doc = Document(str(docx_path))

    # 1. Find the boundaries of the References section
    ref_start_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if REF_HEADER_PATTERN.match(p.text):
            ref_start_idx = i
            break

    if ref_start_idx == -1:
        print("❌ ERROR: Could not locate 'References' header in the document.")
        sys.exit(1)

    ref_end_idx = len(doc.paragraphs)
    for i in range(ref_start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if text and POST_REF_PATTERN.match(text):
            ref_end_idx = i
            break

    # 2. Iterate through the references and append DOIs
    added_count = 0
    already_had_count = 0

    for i in range(ref_start_idx + 1, ref_end_idx):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        # Skip empty lines or very short fragments
        if len(text) < 20: continue

        # Check if a DOI is already present in this paragraph
        if 'doi.org' in text.lower() or 'doi:' in text.lower():
            already_had_count += 1
            continue

        # Fuzzy match the paragraph text against our BibTeX titles
        best_match_doi = None
        best_score = 85  # Minimum strictness threshold
        
        para_clean = clean_for_match(text)
        for bib_title, doi in doi_map.items():
            # partial_ratio is perfect here because the title is just a substring of the full reference paragraph
            score = fuzz.partial_ratio(bib_title, para_clean)
            if score > best_score:
                best_score = score
                best_match_doi = doi

        if best_match_doi:
            # Gather the font styling of the existing text to blend in perfectly
            base_font_name = None
            base_font_size = None
            base_style = None

            original_runs = [r for r in para.runs if r.text.strip()]
            if original_runs:
                reference_run = original_runs[-1]
                base_font_name = reference_run.font.name
                base_font_size = reference_run.font.size
                base_style = reference_run.style

            # Create the new runs
            new_runs = []
            if not text.endswith('.'):
                new_runs.append(para.add_run('.'))
            
            new_runs.append(para.add_run(f" https://doi.org/{best_match_doi}"))
            
            # Apply the captured formatting to the new runs
            for run in new_runs:
                if base_style:
                    run.style = base_style
                if base_font_name:
                    run.font.name = base_font_name
                if base_font_size:
                    run.font.size = base_font_size
                    
            added_count += 1

    # 3. Save the patched draft
    doc.save(str(output_path))
    print(f"\nSuccess! Saved to {output_path.name}")
    print(f" -> Found {already_had_count} references that already had DOIs.")
    print(f" -> Dynamically matched and injected {added_count} missing DOIs.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Appends DOIs to the References section of an intermediate draft.")
    parser.add_argument("bib", type=Path, help="The verified .bib file containing the DOIs")
    parser.add_argument("doc", type=Path, help="The intermediate .docx file")
    args = parser.parse_args()

    if not args.bib.exists():
        print(f"❌ ERROR: BibTeX file '{args.bib}' not found.")
        sys.exit(1)
    if not args.doc.exists():
        print(f"❌ ERROR: Document '{args.doc}' not found.")
        sys.exit(1)

    output = args.doc.with_name(f"{args.doc.stem}_with_DOIs.docx")
    process_document(args.bib, args.doc, output)
