import sys
import os
import re
import argparse
import html
import warnings
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Match

# --- Suppress harmless citeproc-py schema validation warnings ---
warnings.filterwarnings("ignore", category=UserWarning)

# --- MONKEY PATCH FOR PYPARSING/BIBTEXPARSER COMPATIBILITY ---
import pyparsing
if not hasattr(pyparsing, 'DelimitedList'):
    if hasattr(pyparsing, 'delimited_list'): setattr(pyparsing, 'DelimitedList', pyparsing.delimited_list)
    elif hasattr(pyparsing, 'delimitedList'): setattr(pyparsing, 'DelimitedList', pyparsing.delimitedList)

# --- MONKEY PATCHES FOR CITEPROC-PY ENGINE CRASHES ---
try:
    # Fix 1: UnboundLocalError during author name parsing (comma lists without 'and')
    import citeproc.source.bibtex.bibtex as citeproc_bibtex
    _orig_parse_name = citeproc_bibtex.parse_name
    def _robust_parse_name(name):
        try:
            return _orig_parse_name(name)
        except UnboundLocalError:
            clean_name = name.replace(',', ' ')
            try:
                return _orig_parse_name(clean_name)
            except Exception:
                return (None, None, name, None)
        except Exception:
            return (None, None, name, None)
    citeproc_bibtex.parse_name = _robust_parse_name
except Exception:
    pass

try:
    # Fix 2: UnboundLocalError during complex page range formatting (e.g. 3587.e29)
    import citeproc.model as cp_model
    # Dynamically find and patch any class containing '_format_last_page' (e.g., cp_model.Number)
    for name, obj in vars(cp_model).items():
        if isinstance(obj, type) and hasattr(obj, '_format_last_page'):
            _orig_format = getattr(obj, '_format_last_page')
            
            def make_robust_formatter(orig_func):
                def _robust_format_last_page(self, first, last, *args, **kwargs):
                    try:
                        return orig_func(self, first, last, *args, **kwargs)
                    except UnboundLocalError:
                        return last  # Fallback to raw page number
                    except Exception:
                        return last
                return _robust_format_last_page
                
            setattr(obj, '_format_last_page', make_robust_formatter(_orig_format))
except Exception:
    pass


# --- CITEPROC IMPORTS ---
from citeproc import CitationStylesStyle, CitationStylesBibliography
from citeproc import Citation, CitationItem
from citeproc import formatter
from citeproc.source.bibtex import BibTeX
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from rapidfuzz import fuzz

# --- INTELLIGENCE DICTIONARIES ---
AA_LIST_3 = "Ala|Arg|Asn|Asp|Cys|Gln|Glu|Gly|His|Ile|Leu|Lys|Met|Phe|Pro|Ser|Thr|Trp|Tyr|Val"
AA_PATTERN_3 = re.compile(rf'\b(?:{AA_LIST_3})\s*-?\s*$', re.IGNORECASE)
AA_PATTERN_1 = re.compile(r'\b[ARNDCQEGHILKMFPSTWY]-?$')
MATH_UNIT_LIST = ["CV", "R", "r", "m", "cm", "mm", "µm", "um", "nm", "km", "kg", "x", "y", "z", "p", "n", "k", "v", "V", "D", "Ca", "Mg", "Na", "K", "Cl", "Fe", "Zn", "Cu", "O", "H", "N", "C", "P", "S", "M", "χ", "Χ"]
MATH_UNIT_PATTERN = re.compile(rf'\b(?:{"|".join(MATH_UNIT_LIST)})\s*$') 
# A trailing sign is allowed between the mantissa and the exponent: manuscripts vary
# in whether the minus of "4x10^-132" sits inside the superscript run or before it, and
# when it sits outside, a bare-digit exponent would otherwise be read as a citation.
_SIGNS = '-‐‑‒–—―−+'   # hyphen, dashes, true minus, plus
POWER_PATTERN = re.compile(rf'(?:10|x10|×10|\*10)\s*[{_SIGNS}]?\s*$')
IGNORE_PREFIXES = re.compile(r'(?i:\b(?:fig(?:ure)?|eq(?:uation)?|tbl|table|section|sec|step)\s*\.?\s*)$')

PROSE_STOP_WORDS = {'the', 'is', 'are', 'was', 'were', 'that', 'this', 'to', 'for', 'with', 
                    'in', 'on', 'by', 'an', 'we', 'our', 'as', 'it', 'can', 'be', 'has', 'have', 'of', 'and', 'from', 'which'}

REF_HEADER_PATTERN = re.compile(r'^\s*(?:[0-9]+\.?\s*)?(?:REFERENCES|BIBLIOGRAPHY|LITERATURE CITED|WORKS CITED)\s*$', re.IGNORECASE)
POST_REF_PATTERN = re.compile(r'^\s*(?:Tables?|Figures?|Figure Legends?|Supplementary.*?|Appendices|Data Availability|Acknowledgements?|Author Contributions?|Funding|Conflict(?:s)? of Interest|Competing Interests?|(?:Table|Figure|Fig\.?)\s*\d+.*)$', re.IGNORECASE)

def _first_author_surname(entry) -> str:
    """Lowercased surname of an entry's first author, for author-year matching.

    citeproc-py normalises the BibTeX `author` field into a list of Name mappings
    ({'given': ..., 'family': ...}), so string-splitting the field -- which is what
    raw BibTeX would need -- yields nonsense on a citeproc Reference. Both shapes are
    accepted here so the map is built whatever the bib source hands over.
    """
    authors = entry.get('author') or entry.get('editor') or ''

    if isinstance(authors, (list, tuple)):
        if not authors:
            return ''
        first = authors[0]
        if isinstance(first, dict):
            # 'literal' covers institutional authors, which have no family name
            name = first.get('family') or first.get('literal') or ''
        else:
            name = (getattr(first, 'family', None) or getattr(first, 'literal', None)
                    or str(first))
        return str(name).strip().strip('{}').lower()

    # raw BibTeX string: "Family, Given and Family, Given and ..."
    block = str(authors).split(' and ')[0].split(',')[0].split()
    return block[-1].strip('{}').lower() if block else ''


def _issued_year(entry) -> str:
    """A 4-digit year for an entry, as a string, or '' if there isn't one.

    citeproc-py folds BibTeX `year` into `issued`, so entry.get('year') is always
    None on a Reference; reading only 'year' leaves the author-year map empty.
    """
    for source in (entry.get('issued'), entry.get('year'), entry.get('date')):
        if not source:
            continue
        if isinstance(source, dict):
            # 'begin' covers a DateRange, whose year lives on its start date
            candidate = source.get('year') or (source.get('begin') or {}).get('year', '')
        else:
            candidate = source
        match = re.search(r'((?:19|20)\d{2})', str(candidate))
        if match:
            return match.group(1)
    return ''


# --- CITATION MANAGER CLASS ---
class CSLCitationManager:
    def __init__(self, bib_file: Path, csl_file: Path):
        self.bib_file = bib_file
        self.csl_file = csl_file
        self.update_count: int = 0
        
        print(f"Loading Bibliography Data and CSL style ({csl_file.name})...")
        
        self.bib_source = BibTeX(str(self.bib_file), encoding='utf-8')
        self.bib_style = CitationStylesStyle(str(self.csl_file))
        
        if getattr(self.bib_style.root, 'citation', None) is None:
            parent_link = None
            try:
                with open(self.csl_file, 'r', encoding='utf-8') as f:
                    match = re.search(r'<link\s+rel="independent-parent"\s+href="([^"]+)"', f.read())
                    if match: parent_link = match.group(1)
            except Exception: pass
            
            print(f"\n❌ ERROR: '{csl_file.name}' is a 'dependent' CSL style. It does not contain formatting rules.")
            print(f"   citeproc-py requires the full independent parent style to format citations.")
            if parent_link:
                print(f"   👉 Please download the parent style instead: {parent_link.split('/')[-1]}.csl")
            sys.exit(1)
            
        self.bibliography = CitationStylesBibliography(self.bib_style, self.bib_source, formatter.html)
        
        self.is_superscript_style = False
        try:
            with open(self.csl_file, 'r', encoding='utf-8') as f:
                csl_text = f.read()
                if re.search(r'vertical-align\s*=\s*[\'"]sup[\'"]', csl_text, re.IGNORECASE):
                    self.is_superscript_style = True
                elif any(x in str(self.csl_file).lower() for x in ['cell', 'nature', 'lancet', 'science']):
                    self.is_superscript_style = True
        except Exception: pass

        self.id_map: Dict[int, str] = {}
        self.ay_map: List[Dict[str, str]] = []
        
        for key, entry in self.bib_source.items():
            match = re.search(r'(\d+)', key)
            if match:
                ref_num = int(match.group(1))
                self.id_map[ref_num] = key
                
            first_author = _first_author_surname(entry)
            year = _issued_year(entry)
            if first_author and year:
                self.ay_map.append({'key': key, 'author': first_author, 'year': year})

    def get_in_text_citation(self, keys: List[str]) -> str:
        if not keys: return f"[!!MISSING!!]"

        citation_items = [CitationItem(k) for k in keys]
        citation = Citation(citation_items)
        self.bibliography.register(citation)
        
        formatted_cite = self.bibliography.cite(citation, lambda item: None)
        formatted_str = html.unescape(str(formatted_cite)).replace('\u200b', '').replace('\u200c', '').strip()
        clean_text = re.sub(r'<[^>]+>', '', formatted_str).strip()
        
        nums_raw = re.findall(r'\d+', clean_text)
        alpha_chars = re.sub(r'[^A-Za-z]', '', clean_text)
        is_numeric_style = bool(nums_raw) and len(alpha_chars) < 3
        
        if is_numeric_style:
            nums = sorted(list(set(int(n) for n in nums_raw)))
            ranges, start, prev = [], nums[0], nums[0]
            for n in nums[1:]:
                if n == prev + 1: prev = n
                else:
                    ranges.append(str(start) if prev == start else (f"{start}, {prev}" if prev == start + 1 else f"{start}–{prev}"))
                    start = prev = n
            ranges.append(str(start) if prev == start else (f"{start}, {prev}" if prev == start + 1 else f"{start}–{prev}"))
            
            collapsed = ", ".join(ranges)
            
            if self.is_superscript_style: return f"<sup>{collapsed}</sup>"
            else:
                prefix = clean_text[0] if clean_text and clean_text[0] in '[(' else '['
                suffix = clean_text[-1] if clean_text and clean_text[-1] in '])' else ']'
                return f"{prefix}{collapsed}{suffix}"
                
        if self.is_superscript_style and not ('<sup' in formatted_str.lower()):
            formatted_str = re.sub(r'^([\[\(]?)(.*?)([\]\)]?)$', r'\2', clean_text)
            return f"<sup>{formatted_str}</sup>"
            
        return formatted_str

# --- HELPER & DOCX PROCESSING ---

def iter_block_items(doc):
    for child in doc.element.body:
        if isinstance(child, CT_P): yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl): yield Table(child, doc)

def replace_text_preserve_formatting(para: Paragraph, pattern: re.Pattern, callback):
    text = para.text
    matches = list(pattern.finditer(text))
    if not matches: return

    replacements = [callback(m) for m in matches]
    run_map = []
    for r_idx, run in enumerate(para.runs):
        for c_idx in range(len(run.text)): run_map.append((r_idx, c_idx))
    
    if len(run_map) != len(text):
        new_text = text
        for match, rep in zip(reversed(matches), reversed(replacements)):
            start, end = match.span()
            new_text = new_text[:start] + rep + new_text[end:]
        para.text = new_text
        return

    for match, rep in zip(reversed(matches), reversed(replacements)):
        start, end = match.span()
        if rep == match.group(0): continue
        
        start_r_idx, start_c_idx = run_map[start]
        end_r_idx, end_c_idx = run_map[end - 1]
        
        if start_r_idx == end_r_idx:
            run = para.runs[start_r_idx]
            run.text = run.text[:start_c_idx] + rep + run.text[end_c_idx + 1:]
        else:
            run_start = para.runs[start_r_idx]
            run_start.text = run_start.text[:start_c_idx] + rep
            for r_idx in range(start_r_idx + 1, end_r_idx): para.runs[r_idx].text = ""
            run_end = para.runs[end_r_idx]
            run_end.text = run_end.text[end_c_idx + 1:]

def apply_html_formatting_to_runs(para: Paragraph):
    tag_pattern = re.compile(r'(</?(?:i|em|b|strong|sup|sub|span)[^>]*>)', re.IGNORECASE)
    runs = list(para.runs)
    
    for run in runs:
        if not run.text or '<' not in run.text: continue
        parts = tag_pattern.split(run.text)
        if len(parts) == 1: continue
            
        is_i = run.font.italic
        is_b = run.font.bold
        is_sup = run.font.superscript
        is_sub = run.font.subscript
        is_sc = run.font.small_caps
        font_name = run.font.name
        
        parent = run._element.getparent()
        idx = parent.index(run._element)
        
        for part in parts:
            if not part: continue
            part_lower = part.lower()
            
            if part_lower.startswith('<') and part_lower.endswith('>'):
                if part_lower.startswith('<i') or part_lower.startswith('<em'): is_i = True
                elif part_lower.startswith('</i') or part_lower.startswith('</em'): is_i = False
                elif part_lower.startswith('<b') or part_lower.startswith('<strong'): is_b = True
                elif part_lower.startswith('</b') or part_lower.startswith('</strong'): is_b = False
                elif part_lower.startswith('<sup'): is_sup = True
                elif part_lower.startswith('</sup'): is_sup = False
                elif part_lower.startswith('<sub'): is_sub = True
                elif part_lower.startswith('</sub'): is_sub = False
                elif 'small-caps' in part_lower and not part_lower.startswith('</'): is_sc = True
                elif part_lower.startswith('</span'): is_sc = False
            else:
                new_run = para.add_run(part)
                if run.style: 
                    new_run.style = run.style
                    
                if is_i is not None: new_run.font.italic = is_i
                if is_b is not None: new_run.font.bold = is_b
                if is_sup is not None: new_run.font.superscript = is_sup
                if is_sub is not None: new_run.font.subscript = is_sub
                if is_sc is not None: new_run.font.small_caps = is_sc
                if font_name: new_run.font.name = font_name
                
                parent.insert(idx, new_run._element)
                idx += 1
                
        parent.remove(run._element)

def process_paragraph_content(para: Paragraph, manager: CSLCitationManager, citation_pattern: re.Pattern, in_main_body: bool):
    preceding_text = ""
    for run in para.runs:
        text = run.text.strip()
        if in_main_body and run.font.superscript and re.match(r'^[\d,\s\-–]+$', text):
            is_math_power = bool(POWER_PATTERN.search(preceding_text)) and text.isdigit()
            if not (AA_PATTERN_3.search(preceding_text) or AA_PATTERN_1.search(preceding_text) or MATH_UNIT_PATTERN.search(preceding_text) or IGNORE_PREFIXES.search(preceding_text) or is_math_power):
                is_valid = True
                for part in text.replace('–', '-').split(','):
                    if '-' in part:
                        b = part.split('-')
                        if not (len(b) == 2 and b[0].strip().isdigit() and b[1].strip().isdigit()): is_valid = False
                    elif not part.strip().isdigit(): is_valid = False
                if is_valid:
                    run.font.superscript = False
                    run.text = f"[{text}]"
                    
        preceding_text += run.text
    
    artifact_pattern = re.compile(r'(?:geometry|ref|source)\.(\d+)', re.IGNORECASE)
    replace_text_preserve_formatting(para, artifact_pattern, lambda m: f"[{m.group(1)}]")

    def replace_callback(match: Match) -> str:
        preceding = para.text[:match.start()]
        if (AA_PATTERN_3.search(preceding) or AA_PATTERN_1.search(preceding) or MATH_UNIT_PATTERN.search(preceding) or IGNORE_PREFIXES.search(preceding)):
            return match.group(0)

        raw_inner = match.group(1).replace('–', '-')
        if match.group(0).startswith('(') and raw_inner.isdigit() and 1900 <= int(raw_inner) <= 2100: return match.group(0)

        oids = []
        for part in raw_inner.split(','):
            part = part.strip()
            if '-' in part:
                bounds = part.split('-')
                if len(bounds) == 2 and bounds[0].strip().isdigit() and bounds[1].strip().isdigit():
                    start, end = int(bounds[0].strip()), int(bounds[1].strip())
                    if start <= end and (end - start) < 50: oids.extend(range(start, end + 1))
                    else: return match.group(0)
                else: return match.group(0)
            else:
                if not part.isdigit(): return match.group(0)
                oids.append(int(part))
        
        valid_keys = [manager.id_map[o] for o in oids if o in manager.id_map]
        if not valid_keys: return match.group(0)
        
        manager.update_count += 1 
        return manager.get_in_text_citation(valid_keys)

    replace_text_preserve_formatting(para, citation_pattern, replace_callback)

    ay_pattern = re.compile(r'\(([A-Za-z][^()]*?(?:19|20)\d{2}[a-z]?)\)')
    def replace_ay_callback(match: Match) -> str:
        raw_inner = match.group(1)
        if '=' in raw_inner or '+' in raw_inner: return match.group(0)
            
        parts, matched_keys, wanted = raw_inner.split(';'), [], 0
        for part in parts:
            # One part may name the same authors for several works: "Lee et al., 2003, 2006".
            # Every year in the part is a citation of its own; matching only the first
            # silently drops the rest while still reporting the citation as resolved.
            years = re.findall(r'(?:19|20)\d{2}', part)
            if not years: continue
            wanted += len(years)
            author_text = re.sub(r'(?:19|20)\d{2}[a-z]?|et al\.?|,|&', '', part).strip().lower()

            for year in years:
                best_match, best_score = None, 80
                for item in manager.ay_map:
                    if item['year'] == year:
                        score = fuzz.partial_ratio(author_text, item['author'])
                        if score > best_score:
                            best_score = score
                            best_match = item['key']

                if best_match: matched_keys.append(best_match)

        # every year named must have resolved, or the citation is left as written rather
        # than rendered with a subset of its works
        if wanted > 0 and len(matched_keys) == wanted:
            manager.update_count += 1
            return manager.get_in_text_citation(matched_keys)
        return match.group(0)

    replace_text_preserve_formatting(para, ay_pattern, replace_ay_callback)
    apply_html_formatting_to_runs(para)

def write_rich_bibliography_entry(doc: Document, html_text: str, main_font: Optional[str], insert_cursor: Optional[Paragraph] = None):
    p = insert_cursor.insert_paragraph_before() if insert_cursor is not None else doc.add_paragraph()
    html_text = html_text.replace('.. ', '. ').replace('..<', '.<')
    html_text = html_text.replace('</div><div class="csl-right-inline">', '</div><div class="csl-right-inline">' + chr(160))
    html_text = re.sub(r'^((?:<[^>]+>|\s)*)(\[\d+\]|\d+\.)\s*(<[^>]+>)?\s*([A-Za-z])', r'\1\2' + chr(160) + r'\3\4', html_text)
    
    is_bold = is_italic = is_smallcaps = is_sup = is_sub = False
    tokens = re.split(r'(<[^>]+>)', html_text)
    
    for token in tokens:
        if not token: continue
        token_lower = token.lower()
        if token_lower.startswith('<'):
            if token_lower.startswith('<b') or token_lower.startswith('<strong'): is_bold = True
            elif token_lower.startswith('</b') or token_lower.startswith('</strong'): is_bold = False
            elif token_lower.startswith('<i') or token_lower.startswith('<em'): is_italic = True
            elif token_lower.startswith('</i') or token_lower.startswith('</em'): is_italic = False
            elif token_lower.startswith('<sup'): is_sup = True
            elif token_lower.startswith('</sup'): is_sup = False
            elif token_lower.startswith('<sub'): is_sub = True
            elif token_lower.startswith('</sub'): is_sub = False
            elif 'small-caps' in token_lower and not token_lower.startswith('</'): is_smallcaps = True
            elif token_lower.startswith('</span'): is_smallcaps = False
            continue
        
        text_content = html.unescape(token)
        if text_content:
            run = p.add_run(text_content)
            run.bold, run.italic = is_bold, is_italic
            if is_smallcaps: run.font.small_caps = True
            if is_sup: run.font.superscript = True
            if is_sub: run.font.subscript = True
            if main_font: run.font.name = main_font

def process_document(docx_path: Path, output_path: Path, manager: CSLCitationManager):
    print(f"\nProcessing document: {docx_path.name}")
    doc = Document(str(docx_path))
    citation_pattern = re.compile(r'[\[\(]([\d\s,\-–]+)[\]\)]')

    main_font = None
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name: main_font = r.font.name; break
        if main_font: break

    ref_header_element = None
    for p in doc.paragraphs:
        if REF_HEADER_PATTERN.match(p.text): ref_header_element = p._element; break

    in_main_body = False 
    block_counter = 0

    for block in iter_block_items(doc):
        block_counter += 1
        if isinstance(block, Paragraph):
            if ref_header_element is not None and block._element == ref_header_element: break 
            if not in_main_body:
                text_clean = block.text.strip().lower()
                if text_clean in ['abstract', 'introduction', 'background', 'summary', 'methods', 'results']: in_main_body = True
                else:
                    words = re.findall(r'\b[a-z]+\b', text_clean)
                    if len(words) >= 25 and len([w for w in words if w in PROSE_STOP_WORDS]) >= 5: in_main_body = True
                if block_counter > 25 and not in_main_body: in_main_body = True

            process_paragraph_content(block, manager, citation_pattern, in_main_body)
            
        elif isinstance(block, Table):
            if not in_main_body: continue 
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_paragraph_content(para, manager, citation_pattern, in_main_body)

    if manager.update_count == 0:
        print("\n⚠️  WARNING: 0 in-text citations were matched to the bibliography.")
        print("   Skipping reference-list rebuild to avoid DELETING the existing references.")
        print("   (Check that the manuscript's citation style matches the CSL; for narrative")
        print("    author-year citations use arm-add-dois instead of a full CSL rebuild.)")
        doc.save(str(output_path))
        print(f"Saved (references left intact) to {output_path.name}")
        print(" -> Tracked and dynamically updated 0 in-text citations.")
        return

    ref_header_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p._element == ref_header_element: ref_header_index = i; break
            
    insert_cursor = None
    if ref_header_index != -1:
        post_ref_index = -1
        for i in range(ref_header_index + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text.strip()
            if text and POST_REF_PATTERN.match(text):
                post_ref_index = i; break
                
        if post_ref_index != -1:
            paragraphs_to_remove = doc.paragraphs[ref_header_index:post_ref_index]
            insert_cursor = doc.paragraphs[post_ref_index]
            insert_cursor.paragraph_format.page_break_before = True
        else:
            paragraphs_to_remove = doc.paragraphs[ref_header_index:]
            
        for p in paragraphs_to_remove:
            parent = p._element.getparent()
            if parent is not None: parent.remove(p._element)

    p = insert_cursor.insert_paragraph_before() if insert_cursor is not None else doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    run = p.add_run('References')
    run.bold = True
    if main_font: run.font.name = main_font

    print(f"   -> Rebuilding Rich-Text Bibliography via CSL...")
    manager.bibliography.sort()
    for entry in manager.bibliography.bibliography():
        write_rich_bibliography_entry(doc, str(entry), main_font, insert_cursor)

    doc.save(str(output_path))
    print(f"Success! Saved to {output_path.name}")
    print(f" -> Tracked and dynamically updated {manager.update_count} in-text citations.")

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("bib", type=Path, help="Verified .bib file")
    parser.add_argument("doc", type=Path, help="Input .docx file")
    parser.add_argument("--csl", type=Path, required=True, help="Path to the CSL style file")
    args = parser.parse_args()
    
    csl_path = args.csl
    default_csl_dir = Path("~/citation_styles").expanduser()
    
    if not csl_path.exists():
        alt_path = default_csl_dir / csl_path.name
        if alt_path.exists(): csl_path = alt_path
        elif not csl_path.suffix == '.csl':
            alt_path_ext = default_csl_dir / f"{csl_path.name}.csl"
            if alt_path_ext.exists(): csl_path = alt_path_ext
                
    if not csl_path.exists():
        print(f"❌ ERROR: CSL file '{args.csl}' not found locally or in {default_csl_dir}.")
        sys.exit(1)
        
    output = args.doc.with_name(f"{args.doc.stem}_final_{csl_path.stem}.docx")
    mgr = CSLCitationManager(args.bib, csl_path)
    process_document(args.doc, output, mgr)

if __name__ == "__main__":
    main()
