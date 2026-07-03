#!/usr/bin/env python3
"""
arm-diff: Compare two scientific manuscript .docx files.
Generates an annotated third document marking additions (green highlight) and
deletions (red strikethrough), with section-aware and citation-aware filtering.
"""

import re
import sys
import difflib
import argparse
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Reuse section-boundary and math-protection patterns from ref_management ---
try:
    from ref_management.apply_citations import (
        REF_HEADER_PATTERN, POST_REF_PATTERN, PROSE_STOP_WORDS,
        MATH_UNIT_PATTERN, AA_PATTERN_3, AA_PATTERN_1, POWER_PATTERN,
    )
except ImportError:
    REF_HEADER_PATTERN = re.compile(
        r'^\s*(?:[0-9]+\.?\s*)?'
        r'(?:REFERENCES?|BIBLIOGRAPHY|LITERATURE CITED|WORKS CITED)\s*$',
        re.IGNORECASE,
    )
    POST_REF_PATTERN = re.compile(
        r'^\s*(?:Tables?|Figures?|Figure Legends?|Supplementary.*?|Appendices|'
        r'Data Availability|Acknowledgements?|Author Contributions?|Funding|'
        r'Conflict(?:s)? of Interest|Competing Interests?|'
        r'(?:Table|Figure|Fig\.?)\s*\d+.*)\s*$',
        re.IGNORECASE,
    )
    PROSE_STOP_WORDS = {
        'the', 'is', 'are', 'was', 'were', 'that', 'this', 'to', 'for',
        'with', 'in', 'on', 'by', 'an', 'we', 'our', 'as', 'it', 'can',
        'be', 'has', 'have', 'of', 'and', 'from', 'which',
    }
    MATH_UNIT_PATTERN = re.compile(
        r'\b(?:CV|R|r|m|cm|mm|µm|um|nm|km|kg|x|y|z|p|n|k|v|V|D|'
        r'Ca|Mg|Na|K|Cl|Fe|Zn|Cu|O|H|N|C|P|S|M|χ|Χ)\s*$'
    )
    AA_PATTERN_3 = re.compile(
        r'\b(?:Ala|Arg|Asn|Asp|Cys|Gln|Glu|Gly|His|Ile|'
        r'Leu|Lys|Met|Phe|Pro|Ser|Thr|Trp|Tyr|Val)\s*-?\s*$',
        re.IGNORECASE,
    )
    AA_PATTERN_1 = re.compile(r'\b[ARNDCQEGHILKMFPSTWY]-?$')
    POWER_PATTERN = re.compile(r'(?:10|x10|×10|\*10)\s*$')

# OOXML namespace map
_NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# In-text citation patterns
NUMERIC_CITE = re.compile(r'[\[\(]([\d\s,\-–]+)[\]\)]')
AY_CITE      = re.compile(r'\(([A-Za-z][^()]*?(?:19|20)\d{2}[a-z]?)\)')

GREEN  = RGBColor(0,   150,  60)
RED    = RGBColor(200,   0,   0)


# ---------------------------------------------------------------------------
# Document extraction
# ---------------------------------------------------------------------------

def _para_text_excluding_deletions(p_element):
    """Return visible text of a paragraph, excluding w:del tracked deletions."""
    _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    _W_T   = f'{{{_W}}}t'
    _W_DEL = f'{{{_W}}}del'
    result = []
    for node in p_element.iter(_W_T):
        parent = node.getparent()
        in_del = False
        while parent is not None and parent != p_element:
            if parent.tag == _W_DEL:
                in_del = True
                break
            parent = parent.getparent()
        if not in_del and node.text:
            result.append(node.text)
    return "".join(result).strip()


def _is_superscript_numeric_citation(run, preceding_text):
    """Return True if this run looks like a superscript citation number."""
    text = run.text.strip()
    if not (run.font.superscript and re.match(r'^[\d,\s\-–]+$', text)):
        return False
    if (AA_PATTERN_3.search(preceding_text) or AA_PATTERN_1.search(preceding_text)
            or MATH_UNIT_PATTERN.search(preceding_text)
            or POWER_PATTERN.search(preceding_text)):
        return False
    return True


def extract_paragraphs(docx_path, mode, cite_style=None):
    """
    Extract non-empty paragraph texts from a .docx file.

    mode:
        'full'               – entire document
        'main'               – stop at the References section header
        'main-no-citations'  – same as 'main' but in-text citations stripped

    Returns list of strings (one per non-empty paragraph).
    """
    doc = Document(str(docx_path))
    texts = []

    for p in doc.paragraphs:
        raw = _para_text_excluding_deletions(p._element)

        # Normalise superscript citation numbers → [n] so they strip cleanly
        reconstructed = []
        prev = ""
        for run in p.runs:
            if _is_superscript_numeric_citation(run, prev):
                reconstructed.append(f"[{run.text.strip()}]")
            else:
                reconstructed.append(run.text)
            prev = run.text

        # Use normalised text if substantially the same length, else keep raw
        norm = "".join(reconstructed).strip()
        texts.append(norm if abs(len(norm) - len(raw)) < 5 else raw)

    # Find References boundary
    cutoff = len(texts)
    if mode in ('main', 'main-no-citations'):
        for i, t in enumerate(texts):
            if REF_HEADER_PATTERN.match(t):
                cutoff = i
                break

    result = [t for t in texts[:cutoff] if t.strip()]

    if mode == 'main-no-citations' and cite_style:
        result = [_strip_citations(t, cite_style) for t in result]
        result = [t for t in result if t.strip()]

    return result


def detect_citation_style(para_texts):
    """Auto-detect numeric vs. author-year citation style."""
    sample = para_texts[:60]
    numeric = sum(len(NUMERIC_CITE.findall(p)) for p in sample)
    ay      = sum(len(AY_CITE.findall(p))      for p in sample)
    if numeric > ay:
        return 'numeric'
    if ay > 0:
        return 'author-year'
    return 'unknown'


def _strip_citations(text, style):
    if style == 'numeric':
        text = NUMERIC_CITE.sub('', text)
    elif style == 'author-year':
        text = AY_CITE.sub('', text)
    return re.sub(r'  +', ' ', text).strip()


# ---------------------------------------------------------------------------
# Output document construction
# ---------------------------------------------------------------------------

def _add_highlight(run, color_name='yellow'):
    rPr = run._r.get_or_add_rPr()
    h = OxmlElement('w:highlight')
    h.set(qn('w:val'), color_name)
    rPr.append(h)


def _run(para, text, color=None, strike=False, highlight=False, bold=False):
    if not text:
        return None
    r = para.add_run(text)
    if color:
        r.font.color.rgb = color
    if strike:
        r.font.strike = True
    if bold:
        r.font.bold = True
    if highlight:
        _add_highlight(r, 'yellow')
    return r


def _word_diff_paragraph(para, orig_text, edit_text):
    """Render a word-level diff into a single paragraph."""
    orig_tokens = re.findall(r'\S+|\s+', orig_text)
    edit_tokens = re.findall(r'\S+|\s+', edit_text)
    matcher = difflib.SequenceMatcher(None, orig_tokens, edit_tokens, autojunk=False)
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            _run(para, ''.join(orig_tokens[i1:i2]))
        elif op == 'insert':
            _run(para, ''.join(edit_tokens[j1:j2]), color=GREEN, highlight=True)
        elif op == 'delete':
            _run(para, ''.join(orig_tokens[i1:i2]), color=RED, strike=True)
        elif op == 'replace':
            _run(para, ''.join(orig_tokens[i1:i2]), color=RED,   strike=True)
            _run(para, ''.join(edit_tokens[j1:j2]), color=GREEN, highlight=True)


def build_comparison_doc(orig_paras, edit_paras, output_path):
    doc = Document()

    # Header
    title = doc.add_paragraph()
    _run(title, 'DOCUMENT COMPARISON', bold=True).font.size = Pt(14)

    legend = doc.add_paragraph()
    _run(legend, 'Legend:  ')
    _run(legend, 'ADDITION ', color=GREEN, highlight=True)
    _run(legend, '  |  ')
    _run(legend, 'DELETION', color=RED, strike=True)
    doc.add_paragraph()

    matcher = difflib.SequenceMatcher(None, orig_paras, edit_paras, autojunk=False)

    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'equal':
            for text in edit_paras[j1:j2]:
                doc.add_paragraph(text)

        elif op == 'insert':
            for text in edit_paras[j1:j2]:
                p = doc.add_paragraph()
                _run(p, text, color=GREEN, highlight=True)

        elif op == 'delete':
            for text in orig_paras[i1:i2]:
                p = doc.add_paragraph()
                _run(p, text, color=RED, strike=True)

        elif op == 'replace':
            orig_block = orig_paras[i1:i2]
            edit_block = edit_paras[j1:j2]
            if len(orig_block) == len(edit_block):
                # Pair paragraphs one-to-one for word-level diff
                for o, e in zip(orig_block, edit_block):
                    p = doc.add_paragraph()
                    _word_diff_paragraph(p, o, e)
            else:
                # Unequal block sizes: show deletes then inserts paragraph by paragraph
                for text in orig_block:
                    p = doc.add_paragraph()
                    _run(p, text, color=RED, strike=True)
                for text in edit_block:
                    p = doc.add_paragraph()
                    _run(p, text, color=GREEN, highlight=True)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _choose_mode():
    print('\n=== ARM-DIFF: Scientific Document Comparator ===')
    print('\nComparison scope:')
    print('  1  Main text only          (stop before References)')
    print('  2  Main text, no citations (strip in-text citations from diff)')
    print('  3  Full document           (include References and all sections)')
    choice = input('\nSelect [1/2/3] (default: 1): ').strip() or '1'
    return {'1': 'main', '2': 'main-no-citations', '3': 'full'}.get(choice, 'main')


def main():
    parser = argparse.ArgumentParser(
        description='Compare two scientific .docx manuscripts and generate an annotated diff.'
    )
    parser.add_argument('original', nargs='?', help='Original .docx file')
    parser.add_argument('edited',   nargs='?', help='Edited .docx file')
    parser.add_argument('--mode', choices=['main', 'main-no-citations', 'full'],
                        help='Comparison mode (skips interactive prompt)')
    parser.add_argument('--output', help='Output .docx path')
    args = parser.parse_args()

    orig_path = Path(args.original) if args.original else \
                Path(input('Original document (.docx): ').strip().strip('\'"'))
    edit_path = Path(args.edited)   if args.edited   else \
                Path(input('Edited document (.docx):   ').strip().strip('\'"'))

    for p in (orig_path, edit_path):
        if not p.exists():
            print(f'ERROR: File not found: {p}')
            sys.exit(1)

    mode = args.mode or _choose_mode()

    print('\nAnalysing documents...')
    # Detect citation style from a raw extraction of the original
    raw_texts = extract_paragraphs(orig_path, 'main')
    cite_style = detect_citation_style(raw_texts)
    print(f'  Citation style detected: {cite_style}')

    if mode == 'main-no-citations':
        confirm = input(
            f'  Confirm: strip {cite_style} citations before diffing? [Y/n]: '
        ).strip().lower()
        if confirm == 'n':
            mode = 'main'

    print(f'  Mode: {mode}')

    use_cite = cite_style if mode == 'main-no-citations' else None
    orig_paras = extract_paragraphs(orig_path, mode, use_cite)
    edit_paras = extract_paragraphs(edit_path, mode, use_cite)

    print(f'  Original: {len(orig_paras)} paragraphs')
    print(f'  Edited:   {len(edit_paras)} paragraphs')

    out_path = Path(args.output) if args.output else \
               edit_path.with_name(f'{edit_path.stem}_comparison.docx')

    print(f'\nBuilding comparison document...')
    build_comparison_doc(orig_paras, edit_paras, out_path)
    print(f'Done. Saved to: {out_path}')


if __name__ == '__main__':
    main()
